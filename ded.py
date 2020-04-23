from docx import Document
from docx.shared import Inches
import fpdf
fpdf.SYSTEM_TTFONTS = '/path/to/system/fonts'
from fpdf import FPDF

def print_indoc(str_doc):
    # функция для вывода в pdf и docx
    document = Document()

    document.add_heading('Письмо дедушке морозу', 0)

    p = document.add_paragraph('Новогодний текст: \n ')
    p.add_run(str_doc)
    p.add_run('\n\n Вот мое фото: ').bold = True
    document.add_picture('1.jpg', width=Inches(1.25))

    document.save('new_ded.docx')

    pdf = FPDF(orientation='P', unit='mm', format='A4')
    pdf.add_page()
    pdf.set_font("Courier", size=12)
    pdf.add_font('FreeSans', '', 'FreeSans.ttf', uni=True)
    pdf.set_font("FreeSans",size=12)
    #pdf.cell(10, 10, txt='{}'.format(str_doc),ln=0, align=" ")
    pdf.multi_cell(200,5,txt=str_doc)
    pdf.multi_cell(200, 5, txt='\n\n\n\n\n\nВот моя фотка:')
    pdf.image("1.jpg", x=50, y=80, w=100)
    pdf.output("new_ded.pdf")
    pdf.close()

def letter_generation(name, good_behavior,present=None):

    if good_behavior==True:
        present_print='\n-'+'\n-'.join(present)
        str_print=(f"Дорогой,  {name}, в этом году ты себя хорошо вёл и заслужил подарок из этого списка: {present_print}")
    else:
        str_print=(f"Дорогой , {name}, ты себя не очень хорошо вел в этом году, поэтому остаешься без подарка.")

    print(str_print)
    data_infile=open('ded_moroz.txt','w')
    data_infile.write(str_print)
    data_infile.close()
    print_indoc(str_print)


letter_generation('Kostya',True,['магнитофон', 'кинокамера', 'портсигар отечественный', 'куртка замшевая'])
